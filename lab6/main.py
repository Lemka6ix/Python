import dearpygui.dearpygui as dpg
from storage_devices import calculate_hdd, calculate_ssd, calculate_flash
from openpyxl import Workbook
from docx import Document
from docx.shared import Pt
import os

# Константы
RESULTS = []

def calculate_callback():                                   # Обработчик кнопки расчета
    """Calculate button handler"""
    try:
        data_size = dpg.get_value("data_size")
        hdd_price = dpg.get_value("hdd_price")
        ssd_price = dpg.get_value("ssd_price")
        flash_price = dpg.get_value("flash_price")
        
        # Получаем результаты для каждого устройства
        hdd = calculate_hdd(data_size, hdd_price)
        ssd = calculate_ssd(data_size, ssd_price)
        flash = calculate_flash(data_size, flash_price)
        
        global RESULTS
        RESULTS = [hdd, ssd, flash]
        
        # Находим лучшее устройство по цене за ГБ
        best_device = min(RESULTS, key=lambda x: x['price_per_gb'])
        
        # Обновляем UI
        dpg.set_value("hdd_time", f"{hdd['time']:.2f} sec")
        dpg.set_value("hdd_price_gb", f"{hdd['price_per_gb']:.2f} rub/Gb")
        
        dpg.set_value("ssd_time", f"{ssd['time']:.2f} sec")
        dpg.set_value("ssd_price_gb", f"{ssd['price_per_gb']:.2f} rub/Gb")
        
        dpg.set_value("flash_time", f"{flash['time']:.2f} sec")
        dpg.set_value("flash_price_gb", f"{flash['price_per_gb']:.2f} rub/Gb")
        
        dpg.set_value("best_device", 
                     f"Best device: {best_device['type']}\n"
                     f"Price for Gb: {best_device['price_per_gb']:.2f} rub\n"
                     f"Proccesing time: {best_device['time']:.2f} sec")
        
    except Exception as e:
        dpg.set_value("best_device", f"Error: {str(e)}")

def save_to_excel():
    """Saving results to Excel"""
    if not RESULTS:
        return
        
    wb = Workbook()
    ws = wb.active
    ws.title = "Analyze results"
    
    # Заголовки
    ws.append(["Type of device", "Time (sec)", "Price for Gb (rub)", "Speed (Bbit/sec)"])
    
    # Данные
    for device in RESULTS:
        ws.append([device['type'], device['time'], device['price_per_gb'], device['speed']])
    
    # Сохраняем файл
    filename = "storage_analysis.xlsx"
    wb.save(filename)
    os.startfile(filename)

def save_to_word():
    """Saving results to Word"""
    if not RESULTS:
        return
        
    doc = Document()
    
    # Заголовок
    doc.add_heading('Save device\'s analyze', 0)
    
    # Лучшее устройство
    best_device = min(RESULTS, key=lambda x: x['price_per_gb'])
    doc.add_paragraph(f"Best device: {best_device['type']}", style='Heading 2')
    
    # Таблица с результатами
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Заголовки таблицы
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Type of device'
    hdr_cells[1].text = 'Time (sec)'
    hdr_cells[2].text = 'Price of Gb (rub)'
    hdr_cells[3].text = 'Speed (Mbit/sec)'
    
    # Данные таблицы
    for device in RESULTS:
        row_cells = table.add_row().cells
        row_cells[0].text = device['type']
        row_cells[1].text = f"{device['time']:.2f}"
        row_cells[2].text = f"{device['price_per_gb']:.2f}"
        row_cells[3].text = str(device['speed'])
    
    # Сохраняем файл
    filename = "storage_analysis.docx"
    doc.save(filename)
    os.startfile(filename)

def create_gui():
    """Creating a graphical interface"""
    dpg.create_context()
    dpg.create_viewport(title='Save device\'s analyze', width=800, height=600)
    
    with dpg.window(label="Main window", width=780, height=580):
        # Ввод данных
        with dpg.group(horizontal=True):
            dpg.add_text("Data volume (Gb):")
            dpg.add_input_float(tag="data_size", default_value=100, width=100)
        
        with dpg.group(horizontal=True):
            dpg.add_text("Price of HDD (rub):")
            dpg.add_input_float(tag="hdd_price", default_value=3000, width=100)
        
        with dpg.group(horizontal=True):
            dpg.add_text("Price of SSD (rub):")
            dpg.add_input_float(tag="ssd_price", default_value=5000, width=100)
        
        with dpg.group(horizontal=True):
            dpg.add_text("Price of Flash (rub):")
            dpg.add_input_float(tag="flash_price", default_value=1000, width=100)
        
        dpg.add_button(label="Calculate", callback=calculate_callback)
        
        # Результаты
        with dpg.collapsing_header(label="Results"):
            with dpg.group(horizontal=True):
                dpg.add_text("HDD:")
                dpg.add_text(tag="hdd_time", label="Time: ")
                dpg.add_text(tag="hdd_price_gb", label="Price of Gb: ")
            
            with dpg.group(horizontal=True):
                dpg.add_text("SSD:")
                dpg.add_text(tag="ssd_time", label="Time: ")
                dpg.add_text(tag="ssd_price_gb", label="Price of Gb: ")
            
            with dpg.group(horizontal=True):
                dpg.add_text("Flash:")
                dpg.add_text(tag="flash_time", label="Time: ")
                dpg.add_text(tag="flash_price_gb", label="Price of Gb: ")
            
            dpg.add_text(tag="best_device", label="")
        
        # Кнопки экспорта
        with dpg.group(horizontal=True):
            dpg.add_button(label="Save to Excel", callback=save_to_excel)
            dpg.add_button(label="Save to Word", callback=save_to_word)
    
    dpg.setup_dearpygui()
    dpg.show_viewport()
    dpg.start_dearpygui()
    dpg.destroy_context()

if __name__ == "__main__":
    create_gui()